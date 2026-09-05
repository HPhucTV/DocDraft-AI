import io
import docx
from docx.shared import Mm
from app.services.docx_renderer import render_tiptap_to_docx
from app.schemas.document import ExportConfig


def test_render_tiptap_to_docx_structure():
    """Kiểm thử tính toàn vẹn của tệp Word (.docx) được sinh từ Tiptap AST (TASK-113)."""
    sample_ast = {
        "type": "doc",
        "content": [
            {
                "type": "table",
                "attrs": {
                    "isND30Table": True,
                    "tableType": "header",
                },
                "content": [
                    {
                        "type": "tableRow",
                        "content": [
                            {
                                "type": "tableCell",
                                "attrs": {"colWidth": "40%"},
                                "content": [
                                    {
                                        "type": "paragraph",
                                        "attrs": {"textAlign": "center"},
                                        "content": [
                                            {"type": "text", "text": "UBND TỈNH ĐỒNG NAI", "marks": [{"type": "bold"}]}
                                        ],
                                    }
                                ],
                            },
                            {
                                "type": "tableCell",
                                "attrs": {"colWidth": "60%"},
                                "content": [
                                    {
                                        "type": "paragraph",
                                        "attrs": {"textAlign": "center"},
                                        "content": [
                                            {"type": "text", "text": "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM", "marks": [{"type": "bold"}]}
                                        ],
                                    }
                                ],
                            },
                        ],
                    }
                ],
            },
            {
                "type": "heading",
                "attrs": {"level": 2, "textAlign": "center"},
                "content": [{"type": "text", "text": "TỜ TRÌNH"}],
            },
            {
                "type": "paragraph",
                "attrs": {"textAlign": "left"},
                "content": [
                    {"type": "text", "text": "Kính gửi: "},
                    {"type": "text", "text": "[TÊN NGƯỜI NHẬN]", "marks": [{"type": "highlight"}]},
                    {"type": "text", "text": ", số tiền duyệt "},
                    {"type": "text", "text": "120.000.000 đồng", "marks": [{"type": "bold"}, {"type": "underline"}]},
                ],
            },
        ],
    }

    config = ExportConfig(
        margin_top_mm=20,
        margin_bottom_mm=20,
        margin_left_mm=30,
        margin_right_mm=15,
        font_family="Times New Roman",
        font_size_pt=13,
    )

    buffer = render_tiptap_to_docx(sample_ast, config, "Tờ trình kinh phí")

    assert buffer is not None
    assert buffer.getbuffer().nbytes > 0

    # Đọc lại bằng python-docx để kiểm tra không bị lỗi Corrupted Document
    doc = docx.Document(buffer)

    # 1. Kiểm tra lề trang A4 chuẩn NĐ 30
    section = doc.sections[0]
    assert round(section.top_margin.mm) == 20
    assert round(section.bottom_margin.mm) == 20
    assert round(section.left_margin.mm) == 30
    assert round(section.right_margin.mm) == 15

    # 2. Kiểm tra có 1 bảng biểu
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert len(table.rows) == 1
    assert len(table.columns) == 2

    # 3. Kiểm tra nội dung văn bản
    text_found = False
    for p in doc.paragraphs:
        if "TỜ TRÌNH" in p.text:
            text_found = True
            break

    assert text_found, "Không tìm thấy tiêu đề TỜ TRÌNH trong văn bản"
    print("✓ Kiểm thử render Tiptap sang Word .docx thành công 100%!")


if __name__ == "__main__":
    test_render_tiptap_to_docx_structure()
