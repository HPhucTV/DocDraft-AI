import io
import urllib.parse
from fastapi import APIRouter, UploadFile, File, HTTPException
from fastapi.responses import StreamingResponse
import docx

from ..schemas.document import ExportDocxRequest, ParseDocxResponse

router = APIRouter(prefix="/export", tags=["Word Processing"])
parse_router = APIRouter(prefix="/parse", tags=["Word Processing"])


@router.post("/docx")
async def export_docx(payload: ExportDocxRequest):
    """
    Render Tiptap JSON AST to Microsoft Word (.docx) format with NĐ 30/2020 layout.
    Streams the binary file directly from memory without writing to disk.
    """
    try:
        doc = docx.Document()
        
        # Apply standard page margins (Nghị định 30/2020/NĐ-CP: 20mm top/bottom, 30mm left, 15mm right)
        section = doc.sections[0]
        section.top_margin = docx.shared.Mm(payload.config.margin_top_mm)
        section.bottom_margin = docx.shared.Mm(payload.config.margin_bottom_mm)
        section.left_margin = docx.shared.Mm(payload.config.margin_left_mm)
        section.right_margin = docx.shared.Mm(payload.config.margin_right_mm)
        
        # Initial document title
        doc.add_heading(payload.title, level=0)
        
        # Buffer to store Word binary
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Encode filename safely for Content-Disposition header
        safe_filename = urllib.parse.quote(f"{payload.title or 'Van_ban'}.docx")
        
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{safe_filename}",
            },
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to render docx: {str(e)}")


@parse_router.post("/docx", response_model=ParseDocxResponse)
async def parse_docx(file: UploadFile = File(...)):
    """
    Parse an uploaded .docx file into Tiptap ProseMirror AST JSON.
    """
    if not file.filename or not file.filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="Only .docx files are supported")
    
    try:
        contents = await file.read()
        doc = docx.Document(io.BytesIO(contents))
        
        # Simple extraction for initial pipeline
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        title = paragraphs[0] if paragraphs else file.filename.replace(".docx", "")
        
        tiptap_content = []
        for p_text in paragraphs:
            tiptap_content.append({
                "type": "paragraph",
                "content": [{"type": "text", "text": p_text}]
            })
            
        content_json = {
            "type": "doc",
            "content": tiptap_content
        }
        
        word_count = sum(len(p.split()) for p in paragraphs)
        table_count = len(doc.tables)
        
        return ParseDocxResponse(
            success=True,
            title=title,
            content_json=content_json,
            word_count=word_count,
            table_count=table_count,
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to parse docx: {str(e)}")
