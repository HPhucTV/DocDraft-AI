import io
from typing import Dict, Any, List
import docx
from docx.shared import Mm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

from ..schemas.document import ExportConfig


def remove_table_borders(table) -> None:
    """
    Triệt tiêu toàn bộ đường viền đen của bảng theo chuẩn Nghị định 30/2020/NĐ-CP
    (Khối Quốc hiệu - Tiêu ngữ và Nơi nhận - Chữ ký).
    """
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="none"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'<w:insideH w:val="none"/>'
        f'<w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(tblBorders)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150) -> None:
    """Thiết lập padding lề ô trong bảng."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def apply_run_marks(run, marks: List[Dict[str, Any]], font_family: str, default_font_size: int) -> None:
    """Gán các định dạng chữ (Marks) như bold, italic, underline, strike, highlight."""
    run.font.name = font_family
    run.font.size = Pt(default_font_size)
    run.font.color.rgb = RGBColor(0, 0, 0)

    for mark in marks:
        m_type = mark.get("type")
        if m_type == "bold":
            run.bold = True
        elif m_type == "italic":
            run.italic = True
        elif m_type == "underline":
            run.underline = True
        elif m_type == "strike":
            run.font.strike = True
        elif m_type in ("highlight", "safePlaceholderHighlight"):
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def render_paragraph_node(container, node: Dict[str, Any], config: ExportConfig) -> None:
    """Render một đoạn văn bản (Paragraph) từ Tiptap AST sang Word."""
    p = container.add_paragraph()

    # Căn lề đoạn văn
    attrs = node.get("attrs", {}) or {}
    text_align = attrs.get("textAlign", "left")

    if text_align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif text_align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif text_align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Giãn dòng tiêu chuẩn hành chính 1.3 - 1.4 lines
    p.paragraph_format.line_spacing = 1.3
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)

    # Render từng text inline
    content_list = node.get("content", [])
    for inline in content_list:
        if inline.get("type") == "text":
            text_str = inline.get("text", "")
            run = p.add_run(text_str)
            apply_run_marks(run, inline.get("marks", []), config.font_family, config.font_size_pt)
        elif inline.get("type") == "hardBreak":
            p.add_run("\n")


def render_heading_node(doc: docx.Document, node: Dict[str, Any], config: ExportConfig) -> None:
    """Render tiêu đề văn bản (Heading 1-3)."""
    attrs = node.get("attrs", {}) or {}
    level = attrs.get("level", 2)
    text_align = attrs.get("textAlign", "center")

    p = doc.add_paragraph()
    if text_align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif text_align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)

    # Cỡ chữ tiêu đề theo cấp
    size_pt = config.font_size_pt + 2 if level == 1 else config.font_size_pt + 1

    content_list = node.get("content", [])
    for inline in content_list:
        if inline.get("type") == "text":
            run = p.add_run(inline.get("text", ""))
            run.bold = True
            run.font.name = config.font_family
            run.font.size = Pt(size_pt)
            run.font.color.rgb = RGBColor(0, 0, 0)


def render_table_node(doc: docx.Document, node: Dict[str, Any], config: ExportConfig) -> None:
    """Render bảng biểu và bảng ẩn 2 cột chuẩn Nghị định 30/2020/NĐ-CP."""
    rows = node.get("content", [])
    if not rows:
        return

    num_rows = len(rows)
    num_cols = 0
    for r in rows:
        cols_in_row = len(r.get("content", []))
        if cols_in_row > num_cols:
            num_cols = cols_in_row

    if num_cols == 0:
        return

    table = doc.add_table(rows=num_rows, cols=num_cols)

    # Kiểm tra bảng ẩn 2 cột (ND 30)
    attrs = node.get("attrs", {}) or {}
    is_nd30 = attrs.get("isND30Table", False) or attrs.get("data-nd30-table") == "true" or num_cols == 2
    table_type = attrs.get("tableType", "general")

    if is_nd30:
        remove_table_borders(table)

    # Chiều rộng có thể in: A4 (210mm) - 30mm (trái) - 15mm (phải) = 165mm
    total_avail_mm = 210 - config.margin_left_mm - config.margin_right_mm

    col_widths = []
    if is_nd30 and num_cols == 2:
        if table_type == "header":
            # Tỉ lệ 40% / 60%
            col_widths = [Mm(total_avail_mm * 0.40), Mm(total_avail_mm * 0.60)]
        else:
            # Tỉ lệ 50% / 50%
            col_widths = [Mm(total_avail_mm * 0.50), Mm(total_avail_mm * 0.50)]
    else:
        # Chia đều các cột
        col_widths = [Mm(total_avail_mm / num_cols) for _ in range(num_cols)]

    for row_idx, r_node in enumerate(rows):
        cells = r_node.get("content", [])
        for col_idx, c_node in enumerate(cells):
            if col_idx < num_cols:
                cell = table.cell(row_idx, col_idx)
                cell.width = col_widths[col_idx]
                set_cell_margins(cell, top=80, bottom=80, left=100, right=100)

                # Render nội dung trong ô
                cell_contents = c_node.get("content", [])
                # Xóa đoạn văn mặc định rỗng của cell
                if cell.paragraphs and cell_contents:
                    # Clear run trong paragraph đầu tiên
                    first_p = cell.paragraphs[0]
                    first_p.text = ""

                for p_idx, p_node in enumerate(cell_contents):
                    if p_node.get("type") == "paragraph":
                        if p_idx == 0 and cell.paragraphs:
                            p = cell.paragraphs[0]
                            # Render inline vào first_p
                            p_attrs = p_node.get("attrs", {}) or {}
                            t_align = p_attrs.get("textAlign", "left")
                            if t_align == "center":
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            elif t_align == "right":
                                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            elif t_align == "justify":
                                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            else:
                                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                            p.paragraph_format.line_spacing = 1.25
                            p.paragraph_format.space_after = Pt(1)

                            for inline in p_node.get("content", []):
                                if inline.get("type") == "text":
                                    run = p.add_run(inline.get("text", ""))
                                    apply_run_marks(run, inline.get("marks", []), config.font_family, config.font_size_pt)
                        else:
                            render_paragraph_node(cell, p_node, config)


def render_tiptap_to_docx(content_json: Dict[str, Any], config: ExportConfig, title: str) -> io.BytesIO:
    """
    Bộ động cơ duyệt toàn bộ cây cú pháp Tiptap JSON AST và render tài liệu Word (.docx)
    chuẩn Nghị định 30/2020/NĐ-CP (TASK-113).
    """
    doc = docx.Document()

    # 1. Thiết lập khổ giấy A4 và lề trang quy chuẩn (NĐ 30: 20mm trên/dưới, 30mm trái, 15mm phải)
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(config.margin_top_mm)
    section.bottom_margin = Mm(config.margin_bottom_mm)
    section.left_margin = Mm(config.margin_left_mm)
    section.right_margin = Mm(config.margin_right_mm)

    # 2. Thiết lập Style mặc định Times New Roman
    normal_style = doc.styles["Normal"]
    normal_style.font.name = config.font_family
    normal_style.font.size = Pt(config.font_size_pt)

    # 3. Duyệt cây Tiptap AST
    doc_nodes = content_json.get("content", [])

    for node in doc_nodes:
        n_type = node.get("type")

        if n_type == "paragraph":
            render_paragraph_node(doc, node, config)
        elif n_type == "heading":
            render_heading_node(doc, node, config)
        elif n_type == "table":
            render_table_node(doc, node, config)
        elif n_type == "bulletList":
            for item in node.get("content", []):
                for p_item in item.get("content", []):
                    p = doc.add_paragraph(style="List Bullet")
                    p.paragraph_format.line_spacing = 1.3
                    for inline in p_item.get("content", []):
                        if inline.get("type") == "text":
                            run = p.add_run(inline.get("text", ""))
                            apply_run_marks(run, inline.get("marks", []), config.font_family, config.font_size_pt)
        elif n_type == "orderedList":
            for item in node.get("content", []):
                for p_item in item.get("content", []):
                    p = doc.add_paragraph(style="List Number")
                    p.paragraph_format.line_spacing = 1.3
                    for inline in p_item.get("content", []):
                        if inline.get("type") == "text":
                            run = p.add_run(inline.get("text", ""))
                            apply_run_marks(run, inline.get("marks", []), config.font_family, config.font_size_pt)
        elif n_type == "horizontalRule":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("─────────────────")
            run.font.name = config.font_family
            run.font.size = Pt(10)

    # 4. Xuất ra binary stream trong RAM (Zero-disk streaming)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
